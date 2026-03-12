from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_uv_guide():
    doc = Document()
    
    # Title
    title = doc.add_heading('UV 使用指南', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Introduction
    doc.add_paragraph('简介')
    intro = doc.add_paragraph('UV 是一个用 Rust 编写的高性能 Python 包管理器和安装程序，旨在成为 pip、venv 和 pip-tools 的现代替代品。它具有极快的速度（比 pip 快 10-100 倍）、可靠的依赖解析以及现代化的功能。')
    
    # Installation section
    doc.add_heading('安装', level=1)
    install_para = doc.add_paragraph('UV 可以通过多种方式安装：')
    
    doc.add_paragraph('使用官方安装脚本（推荐）：', style='Heading 2')
    doc.add_paragraph('curl -LsSf https://astral.sh/uv/0.6.12/install.sh | sh', style='Intense Quote')
    doc.add_paragraph('或使用 pip：')
    doc.add_paragraph('pip install uv', style='Intense Quote')
    doc.add_paragraph('Windows PowerShell 用户可以使用：')
    doc.add_paragraph('irm https://astral.sh/uv/0.6.12/install.ps1 | iex', style='Intense Quote')
    
    # Basic Usage section
    doc.add_heading('基本使用', level=1)
    
    doc.add_paragraph('创建新项目', style='Heading 2')
    doc.add_paragraph('uv init my-project')
    doc.add_paragraph('这将创建一个新的项目目录，包含 pyproject.toml 文件。')
    
    doc.add_paragraph('同步依赖', style='Heading 2')
    doc.add_paragraph('uv sync')
    doc.add_paragraph('从 pyproject.toml 安装依赖项到虚拟环境中。')
    
    doc.add_paragraph('添加依赖', style='Heading 2')
    doc.add_paragraph('uv add requests')
    doc.add_paragraph('uv add "django>=4.0"')
    doc.add_paragraph('uv add --dev pytest')
    
    doc.add_paragraph('运行命令', style='Heading 2')
    doc.add_paragraph('uv run python script.py')
    doc.add_paragraph('uv run pytest')
    
    # Virtual Environments section
    doc.add_heading('虚拟环境管理', level=1)
    
    doc.add_paragraph('创建虚拟环境', style='Heading 2')
    doc.add_paragraph('uv venv')
    doc.add_paragraph('uv venv .venv  # 指定目录名称')
    
    doc.add_paragraph('解释 UV 如何管理虚拟环境', style='Heading 2')
    doc.add_paragraph('UV 自动在项目目录中创建 .venv 文件夹（如果不存在）。')
    doc.add_paragraph('使用 UV_RUN_VERBOSE=1 环境变量查看详细输出。')
    
    # Package Management section
    doc.add_heading('包管理', level=1)
    
    doc.add_paragraph('安装包', style='Heading 2')
    doc.add_paragraph('uv pip install requests')
    doc.add_paragraph('uv pip install -r requirements.txt')
    
    doc.add_paragraph('列出已安装包', style='Heading 2')
    doc.add_paragraph('uv pip list')
    
    doc.add_paragraph('卸载包', style='Heading 2')
    doc.add_paragraph('uv pip uninstall requests')
    
    doc.add_paragraph('冻结依赖', style='Heading 2')
    doc.add_paragraph('uv pip freeze > requirements.txt')
    
    # Common Commands section
    doc.add_heading('常用命令参考', level=1)
    
    commands = [
        ('初始化新项目', 'uv init [项目目录]'),
        ('同步/安装依赖', 'uv sync'),
        ('添加依赖', 'uv add <包名>'),
        ('移除依赖', 'uv remove <包名>'),
        ('更新依赖', 'uv pip install --upgrade <包名>'),
        ('运行脚本', 'uv run <命令>'),
        ('创建虚拟环境', 'uv venv [目录]'),
        ('激活虚拟环境', 'source .venv/bin/activate (Linux/Mac) 或 .venv\\Scripts\\activate (Windows)'),
        ('安装包', 'uv pip install <包名>'),
        ('导出依赖', 'uv pip freeze > requirements.txt'),
        ('查看已安装包', 'uv pip list'),
        ('清理缓存', 'uv cache clean'),
        ('更新UV自身', 'uv self update'),
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '功能'
    hdr_cells[1].text = '命令'
    
    for func, cmd in commands:
        row_cells = table.add_row().cells
        row_cells[0].text = func
        row_cells[1].text = cmd
    
    # Advantages section
    doc.add_heading('UV 的优势', level=1)
    advantages = [
        '极快的速度 - 比 pip 快 10-100 倍',
        '基于 Ruff 的成熟解析器 - 准确的依赖解析',
        '支持 pip 的 requirements.txt 格式',
        '完全兼容 pip 和 virtualenv',
        '开箱即用的 Pylance/pyright 支持',
        '现代化的锁文件格式（uv.lock）',
        '支持月球车（mooncake）等等。',
    ]
    
    for advantage in advantages:
        doc.add_paragraph(advantage, style='List Bullet')
    
    # pyproject.toml section
    doc.add_heading('pyproject.toml 配置示例', level=1)
    doc.add_paragraph('一个基本的 pyproject.toml 文件示例：', style='Heading 2')
    
    config = """[project]
name = "my-project"
version = "0.1.0"
dependencies = [
    "requests>=2.28.0",
    "click>=8.0.0",
]

[project.optional-dependencies]
dev = [
    "pytest>=7.0.0",
    "black>=22.0.0",
]

[build-system]
requires = ["hatchling"]
build-backend = "hatchling.build"""
    
    doc.add_paragraph(config, style='Intense Quote')
    
    # Dependency Groups section
    doc.add_heading('依赖组功能', level=1)
    doc.add_paragraph('UV 支持将依赖项分组，更灵活地管理开发依赖。', style='Heading 2')
    doc.add_paragraph('示例：')
    group_example = """uv add --group test pytest pytest-cov
uv add --group lint ruff black
uv add --group docs mkdocs mkdocs-material"""
    doc.add_paragraph(group_example, style='Intense Quote')
    doc.add_paragraph('使用 --all-extras 标志安装所有可选依赖组：')
    doc.add_paragraph('uv sync --all-extras', style='Intense Quote')
    
    # Lock file section
    doc.add_heading('锁文件 (uv.lock)', level=1)
    doc.add_paragraph('UV 自动生成和维护一个锁文件（uv.lock），确保依赖版本的可重现性。')
    doc.add_paragraph('使用 uv sync 时，UV 会检查现有锁文件并仅当需要时更新。')
    doc.add_paragraph('要强制重新生成锁文件：')
    doc.add_paragraph('uv sync --refresh', style='Intense Quote')
    
    # Troubleshooting section
    doc.add_heading('故障排除', level=1)
    
    issues = [
        ('虚拟环境未自动创建', '使用 uv sync --frozen 强制创建，或检查目录权限'),
        ('依赖冲突', '运行 uv sync --no-dev 或 uv pip check 检查冲突'),
        ('速度慢', '确保使用最新版本，可通过 uv self update 更新'),
        ('Windows 激活问题', '使用 .venv\\Scripts\\activate 或直接使用 uv run'),
    ]
    
    table2 = doc.add_table(rows=1, cols=2)
    table2.style = 'Light Grid Accent 1'
    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = '问题'
    hdr_cells2[1].text = '解决方案'
    
    for issue, solution in issues:
        row_cells = table2.add_row().cells
        row_cells[0].text = issue
        row_cells[1].text = solution
    
    # Resources section
    doc.add_heading('更多资源', level=1)
    resources = [
        ('官方文档', 'https://docs.astral.sh/uv/'),
        ('GitHub 仓库', 'https://github.com/astral-sh/uv'),
        ('在线演示', 'uv-docs.astral.sh'),
    ]
    
    for resource, link in resources:
        p = doc.add_paragraph()
        run = p.add_run(f'{resource}: ')
        run.bold = True
        p.add_run(link)
    
    # Save the document
    output_path = 'uv_使用指南.docx'
    doc.save(output_path)
    print(f'文档已保存为: {output_path}')
    return output_path

if __name__ == '__main__':
    create_uv_guide()